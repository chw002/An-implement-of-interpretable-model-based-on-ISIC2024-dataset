"""Render ``dissertation_bristol_fyp.md`` to a .docx.

The markdown source is authored in a fixed structure so we hand-compile it
rather than using a general markdown→docx converter.  Importantly, figures
are **never** auto-appended: the markdown itself marks figure placeholders
via lines like ::

    ![](outputs_scale/figures/fig_metrics_vs_data_size.png)

and the renderer inserts the file in place with the supplied caption.  This
removes the duplicate-caption / appendix-dumping behaviour the previous
version produced.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "dissertation_bristol_fyp.md"
OUT = ROOT / "dissertation_bristol_fyp.docx"


FIG_PATTERN = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


# --------------------------------------------------------------- styling ----
def _set_body_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    for h in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[h]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 3"].font.size = Pt(12)


def _add_paragraph(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0.6)
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold


def _add_code_block(doc: Document, lines: List[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _add_markdown_table(doc: Document, raw_lines: List[str]) -> None:
    rows = [
        [c.strip() for c in line.strip("|").split("|")]
        for line in raw_lines
        if not set(line.strip()) <= set("|-: ")
    ]
    if not rows:
        return
    n_cols = len(rows[0])
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cells = tbl.rows[i].cells
            if j >= len(cells):
                break
            cells[j].text = val
            for para in cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True


def _add_figure(
    doc: Document, path: Path, caption: Optional[str] = None, width_cm: float = 13.0
) -> None:
    if not path.exists():
        _add_paragraph(
            doc,
            f"[Figure not available on this run: {path.relative_to(ROOT) if path.is_absolute() else path}]",
            italic=True,
        )
        if caption:
            _add_paragraph(doc, caption, italic=True, bold=True)
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)


def _inline_clean(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^\*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^\*]+)\*", r"\1", text)
    return text


def _resolve_figure(rel: str) -> Path:
    p = Path(rel)
    return p if p.is_absolute() else (ROOT / p)


# --------------------------------------------------------------- main -------
def render() -> None:
    if not SRC.exists():
        raise FileNotFoundError(f"{SRC} not found — nothing to render")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    _set_body_style(doc)

    # ----- Title page
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("GlassDerm")
    run.font.size = Pt(30)
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "An Image-Only, Fully-Auditable Neuro-Symbolic Pipeline\n"
        "for Skin Lesion Classification on ISIC 2024"
    )
    run.font.size = Pt(16)
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    for text, size, italic in [
        ("Final-Year Research Project", 14, False),
        ("BSc Computer Science", 12, True),
        ("", 12, False),
        ("Department of Computer Science", 12, False),
        ("University of Bristol", 12, True),
        ("", 12, False),
        ("2026", 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.italic = italic

    doc.add_page_break()

    # ----- Parse body
    lines = SRC.read_text().splitlines()
    i = 0
    while i < len(lines) and not lines[i].startswith("# "):
        i += 1

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            i += 1
            continue

        # Headings
        if line.startswith("# Chapter") or line.startswith("# References") or line.startswith("# Appendix"):
            doc.add_page_break()
            doc.add_heading(_inline_clean(line.lstrip("# ").rstrip()), level=1)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(_inline_clean(line.lstrip("# ").rstrip()), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(_inline_clean(line.lstrip("# ").rstrip()), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(_inline_clean(line.lstrip("# ").rstrip()), level=3)
            i += 1
            continue

        # Inline figure marker:  ![caption](path/to/fig.png)
        m = FIG_PATTERN.match(stripped)
        if m:
            caption = _inline_clean(m.group(1)) or None
            fig_path = _resolve_figure(m.group(2))
            _add_figure(doc, fig_path, caption=caption)
            i += 1
            continue

        # Fenced code blocks
        if stripped.startswith("```"):
            block: List[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            _add_code_block(doc, block)
            continue

        # Markdown tables
        if stripped.startswith("|"):
            tbl_lines: List[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            _add_markdown_table(doc, tbl_lines)
            continue

        # Bullet lists
        if stripped.startswith("- "):
            while i < len(lines) and lines[i].strip().startswith("- "):
                p = doc.add_paragraph(
                    _inline_clean(lines[i].strip()[2:]),
                    style="List Bullet",
                )
                p.paragraph_format.space_after = Pt(2)
                i += 1
            continue
        if re.match(r"^\d+\.\s", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                body = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(_inline_clean(body), style="List Number")
                p.paragraph_format.space_after = Pt(2)
                i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        _add_paragraph(doc, _inline_clean(line.strip()))
        i += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    render()
