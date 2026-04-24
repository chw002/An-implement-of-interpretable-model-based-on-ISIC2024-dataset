#!/usr/bin/env python3
"""Render a per-scale ``README_RESULTS.md`` to a .docx.

``scale-experiment`` writes a ``README_RESULTS.md`` into every
``outputs_scale/benign_{budget}/`` directory summarising that budget's run
(main metrics, thresholds, feature audit summary, a couple of case studies).
This script turns one such Markdown report into a Word document.

Usage
-----
    python scripts/make_report_docx.py                             # benign_full
    python scripts/make_report_docx.py --scale benign_40000        # another budget
    python scripts/make_report_docx.py --md path/to/README.md      # explicit path

Intentionally a small, generic Markdown → DOCX converter — it does not
hard-code any numbers, pipelines, or claims about results.  The source
``README_RESULTS.md`` is the single source of truth.
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SCALES_ROOT = ROOT / "outputs_scale"

FIG_PATTERN = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def _inline_clean(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^\*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^\*]+)\*", r"\1", text)
    return text


def _add_code_block(doc: Document, lines: List[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _add_markdown_table(doc: Document, raw_lines: List[str]) -> None:
    rows = [
        [c.strip() for c in line.strip("|").split("|")]
        for line in raw_lines
        if not set(line.strip()) <= set("|-: ")
    ]
    if not rows:
        return
    n_cols = len(rows[0])
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cells = tbl.rows[i].cells
            if j >= len(cells):
                break
            cells[j].text = val
            for para in cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True


def _add_figure(
    doc: Document,
    src_dir: Path,
    rel: str,
    caption: Optional[str],
    width_cm: float = 13.0,
) -> None:
    path = Path(rel)
    resolved = path if path.is_absolute() else (src_dir / path).resolve()
    if not resolved.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Figure not available on this run: {rel}]")
        run.italic = True
        if caption:
            cp = doc.add_paragraph()
            run = cp.add_run(caption)
            run.italic = True
            run.bold = True
        return
    doc.add_picture(str(resolved), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)


def _render(md_path: Path, out_path: Path) -> None:
    src_dir = md_path.parent
    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"GlassDerm — {src_dir.name}")
    run.bold = True
    run.font.size = Pt(18)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"Source: {md_path.relative_to(ROOT)}")
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    doc.add_paragraph()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(_inline_clean(line[2:].rstrip()), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(_inline_clean(line[3:].rstrip()), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(_inline_clean(line[4:].rstrip()), level=3)
            i += 1
            continue

        m = FIG_PATTERN.match(stripped)
        if m:
            _add_figure(doc, src_dir, m.group(2), _inline_clean(m.group(1)) or None)
            i += 1
            continue

        if stripped.startswith("```"):
            block: List[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            _add_code_block(doc, block)
            continue

        if stripped.startswith("|"):
            tbl_lines: List[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            _add_markdown_table(doc, tbl_lines)
            continue

        if stripped.startswith("- "):
            while i < len(lines) and lines[i].strip().startswith("- "):
                p = doc.add_paragraph(
                    _inline_clean(lines[i].strip()[2:]),
                    style="List Bullet",
                )
                p.paragraph_format.space_after = Pt(2)
                i += 1
            continue
        if re.match(r"^\d+\.\s", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                body = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(_inline_clean(body), style="List Number")
                p.paragraph_format.space_after = Pt(2)
                i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.add_run(_inline_clean(line.rstrip()))
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path}")


def _resolve_md(args: argparse.Namespace) -> Path:
    if args.md:
        p = Path(args.md)
        return p if p.is_absolute() else (ROOT / p)
    scale_dir = SCALES_ROOT / args.scale
    return scale_dir / "README_RESULTS.md"


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument(
        "--scale",
        default="benign_full",
        help="Subdirectory name under outputs_scale/ (default: benign_full).",
    )
    parser.add_argument(
        "--md",
        default=None,
        help="Explicit path to a README_RESULTS.md; overrides --scale.",
    )
    parser.add_argument(
        "--out",
        default=None,
        help="Destination .docx path (default: alongside the source MD).",
    )
    args = parser.parse_args(argv)

    md = _resolve_md(args)
    if not md.exists():
        print(f"error: {md} not found — run scale-experiment first.", file=sys.stderr)
        return 2
    out = Path(args.out) if args.out else md.with_suffix(".docx")
    _render(md, out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
